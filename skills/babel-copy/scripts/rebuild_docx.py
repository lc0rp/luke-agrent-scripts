#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from collections import defaultdict
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a DOCX from babel-copy translated blocks")
    parser.add_argument("translated_blocks_json")
    parser.add_argument("--output-docx", required=True)
    return parser.parse_args()


def alignment_for(value: str | None) -> WD_ALIGN_PARAGRAPH:
    if value == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if value == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def configure_document(doc: Document, first_page: dict) -> None:
    section = doc.sections[0]
    section.page_width = Pt(first_page["width"])
    section.page_height = Pt(first_page["height"])
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    section.top_margin = Pt(32)
    section.bottom_margin = Pt(32)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)


def set_paragraph_runs(paragraph, block: dict, text: str) -> None:
    run = paragraph.add_run(text)
    paragraph.alignment = alignment_for(block.get("align"))
    role = block.get("role")
    style = block.get("style", {})
    if role in {"heading", "title", "form_label"} or style.get("bold"):
        run.bold = True
    if style.get("italic"):
        run.italic = True
    font_size = style.get("font_size_hint")
    if font_size:
        run.font.size = Pt(max(8.0, min(14.0, float(font_size))))


def add_block_paragraph(container, block: dict, text: str | None = None):
    paragraph = container.add_paragraph()
    set_paragraph_runs(paragraph, block, text or block.get("translated_text") or block.get("text") or "")
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def width_scale_for_table(page: dict, table: dict, available_width_pt: float) -> float:
    table_width = float(table["bbox"][2] - table["bbox"][0]) or available_width_pt
    return available_width_pt / table_width


def add_signature_picture(cell, asset: dict, cell_bbox: list[float], scale: float) -> None:
    image_path = Path(asset["path"])
    if not image_path.exists():
        return
    asset_bbox = asset.get("bbox", cell_bbox)
    cell_width = max(10.0, cell_bbox[2] - cell_bbox[0])
    asset_width = max(12.0, asset_bbox[2] - asset_bbox[0])
    fit_width = min(cell_width * scale * 0.86, asset_width * scale)

    reference_paragraph = cell.add_paragraph()
    reference_paragraph.paragraph_format.space_before = Pt(max(2.0, min(18.0, (asset_bbox[1] - cell_bbox[1]) * 0.22 * scale)))
    center_ratio = ((asset_bbox[0] + asset_bbox[2]) / 2 - cell_bbox[0]) / cell_width
    if center_ratio <= 0.4:
        reference_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif center_ratio >= 0.6:
        reference_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        reference_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    with Image.open(image_path) as image:
        width_px, height_px = image.size
    if width_px <= 0 or height_px <= 0:
        return
    run = reference_paragraph.add_run()
    run.add_picture(str(image_path), width=Pt(fit_width))


def fill_table_cell(cell, cell_data: dict, blocks_by_id: dict, assets_by_id: dict, scale: float) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    first = True
    ordered_blocks = [blocks_by_id[block_id] for block_id in cell_data.get("block_ids", []) if block_id in blocks_by_id]
    ordered_blocks.sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))
    for block in ordered_blocks:
        paragraph = cell.paragraphs[0] if first and cell.paragraphs else cell.add_paragraph()
        if first:
            paragraph.text = ""
        set_paragraph_runs(paragraph, block, block.get("translated_text") or block.get("text") or "")
        paragraph.paragraph_format.space_after = Pt(2)
        first = False
    for asset_id in cell_data.get("signature_asset_ids", []):
        asset = assets_by_id.get(asset_id)
        if asset:
            add_signature_picture(cell, asset, cell_data["bbox"], scale)


def render_table(doc: Document, page: dict, table: dict, blocks_by_id: dict, assets_by_id: dict) -> None:
    rows = table["rows"]
    columns = table["columns"]
    doc_table = doc.add_table(rows=len(rows) - 1, cols=len(columns) - 1)
    doc_table.style = "Table Grid"
    doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc_table.autofit = False

    available_width = page["width"] - 72
    scale = width_scale_for_table(page, table, available_width)

    for col_index in range(len(columns) - 1):
        width_pt = max(42.0, (columns[col_index + 1] - columns[col_index]) * scale)
        for row in doc_table.rows:
            row.cells[col_index].width = Pt(width_pt)

    cells_by_position = {(cell["row_index"], cell["col_index"]): cell for cell in table["cells"]}
    for row_index in range(len(rows) - 1):
        row_height = max(26.0, (rows[row_index + 1] - rows[row_index]) * scale)
        doc_row = doc_table.rows[row_index]
        doc_row.height = Pt(row_height)
        for col_index in range(len(columns) - 1):
            fill_table_cell(doc_row.cells[col_index], cells_by_position[(row_index, col_index)], blocks_by_id, assets_by_id, scale)


def main() -> int:
    args = parse_args()
    payload = json.loads(Path(args.translated_blocks_json).read_text())
    pages = payload.get("pages", [])
    blocks = payload["blocks"]
    assets = payload.get("assets", [])
    if not pages:
        raise SystemExit("Expected pages metadata in translated blocks payload.")

    blocks_by_id = {block["id"]: block for block in blocks}
    assets_by_id = {asset["id"]: asset for asset in assets}
    blocks_by_page: dict[int, list[dict]] = defaultdict(list)
    for block in blocks:
        blocks_by_page[int(block["page_number"])].append(block)
    for page_blocks in blocks_by_page.values():
        page_blocks.sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))

    doc = Document()
    configure_document(doc, pages[0])

    first_page = True
    for page in pages:
        page_number = int(page["page_number"])
        if not first_page:
            doc.add_page_break()
        first_page = False

        page_blocks = blocks_by_page.get(page_number, [])
        for block in page_blocks:
            if block.get("keep_original"):
                continue
            if block.get("table"):
                continue
            if block.get("role") == "artifact":
                continue
            add_block_paragraph(doc, block)

        for table in sorted(page.get("tables", []), key=lambda item: item["bbox"][1]):
            render_table(doc, page, table, blocks_by_id, assets_by_id)

    output_path = Path(args.output_docx).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
